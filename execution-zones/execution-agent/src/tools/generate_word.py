"""
generate_word tool for Execution Agent (027-slack-file-generation-best-practices).

Produces Word (.docx) documents and stores them in tool_context.invocation_state
for the handler to extract and build file_artifact.
"""

from io import BytesIO
from typing import Any, List

from strands import tool

import file_config as fc
from file_config import sanitize_filename

_DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@tool(context=True)
def generate_word(
    filename: str,
    title: str,
    sections: List[dict],
    tool_context: Any,
) -> str:
    """Word文書 (.docx) を生成します。

    タイトルとセクション（見出し + 本文）で構成されたドキュメントを作成します。

    Args:
        filename: ファイル名（拡張子不要、自動付加）
        title: ドキュメントタイトル
        sections: セクションの配列。各要素は heading, content を持つ。
        tool_context: Strands framework context (injected). Contains invocation_state.

    Returns:
        モデルに返す説明文（日本語）。ファイルは invocation_state に格納される。
    """
    if not filename or not isinstance(filename, str) or not filename.strip():
        return "エラー: filename を指定してください。"
    if not title or not isinstance(title, str):
        return "エラー: title を指定してください。"
    if not sections or not isinstance(sections, list):
        return "エラー: sections は必須です（heading, content を持つオブジェクトの配列）。"

    safe_filename = sanitize_filename(filename.strip(), "docx")
    if not safe_filename or "." not in safe_filename:
        safe_filename = f"{safe_filename}.docx"

    try:
        from docx import Document
    except ImportError:
        return "エラー: python-docx がインストールされていません。"

    doc = Document()
    doc.add_heading(str(title), level=0)

    for section_def in sections:
        if not isinstance(section_def, dict):
            continue
        heading = section_def.get("heading")
        content = section_def.get("content")

        if heading:
            doc.add_heading(str(heading), level=1)
        if content:
            doc.add_paragraph(str(content))

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    file_bytes = buf.getvalue()

    if len(file_bytes) > fc.MAX_OFFICE_FILE_BYTES:
        return (
            f"エラー: ファイルサイズが上限（{fc.MAX_OFFICE_FILE_BYTES} バイト）を超えています。"
            f"現在 {len(file_bytes)} バイトです。"
        )

    invocation_state = getattr(tool_context, "invocation_state", {})
    if isinstance(invocation_state, dict):
        invocation_state["generated_file"] = {
            "file_bytes": file_bytes,
            "file_name": safe_filename,
            "mime_type": _DOCX_MIME,
            "description": f"Word文書「{safe_filename}」を生成しました。",
        }

    return f"ファイル「{safe_filename}」を作成しました。Slack にアップロードされます。"
